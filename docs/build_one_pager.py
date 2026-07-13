"""Fill the Project Showcase One-Pager template with real resume-forge content,
reusing every existing run's formatting (font, size, bold, color) untouched --
only the text is replaced. Character counts are checked against the template's
own original placeholder length as a one-page budget guide.

Kept here as a build record, not a portable script: SRC points at the
original template's path on this machine and won't exist elsewhere. To
regenerate docs/resume-forge-one-pager.docx, update SRC and re-run. One-page
fit was verified empirically (Pages.app export -> pdfplumber page count),
not just by the character-budget heuristic below -- see the commit message
for the iteration history if the layout needs revisiting.
"""
import shutil
from pathlib import Path

import docx
from docx.shared import Inches

SRC = "/Users/vasanth/Documents/vishal/Project_Showcase_One_Pager_Template.docx"
OUT = "/Users/vishal/Career/Projects/github/Resume-Tailor-ATS-Scorer/docs/resume-forge-one-pager.docx"
SCREENSHOT = "/Users/vishal/Career/Projects/github/Resume-Tailor-ATS-Scorer/docs/screenshot.png"

shutil.copyfile(SRC, OUT)
d = docx.Document(OUT)

budget_report = []


def set_run(run, text, budget=None):
    if budget is not None:
        budget_report.append((text[:40], len(text), budget))
    run.text = text


paras = d.paragraphs

# --- Title (fix template's own typo: runs read "howcase" not "Showcase") ---
paras[0].runs[0].text = "resume-forge "
paras[0].runs[1].text = "Project Showcase"

# --- Tagline ---
set_run(
    paras[1].runs[0],
    "LLM-tailored resumes that clear ATS filters — with a mechanical "
    "no-fabrication guarantee.",
    budget=97,
)

# --- Remove the empty spacer paragraph after the tagline; the heading's own
# "space before" already provides visual separation, and every inch counts
# on a strict one-pager. ---
spacer = paras[2]._p
spacer.getparent().remove(spacer)

# --- 1. The Big Picture ---
set_run(
    paras[5].runs[0],
    "ATS software rejects strong candidates for mechanical reasons — "
    "mismatched keywords, broken layouts — and manually re-tailoring for "
    "every posting does not scale.",
    budget=164,
)
set_run(
    paras[7].runs[0],
    "resume-forge parses a resume once, tailors it per posting, renders a "
    "one-page LaTeX PDF, and scores it locally — while a guard blocks any "
    "invented fact.",
    budget=166,
)

# --- 2. Core Features & Capabilities (bullets: [lead-in bold run, desc run]) ---
features = [
    (9, "No-Fabrication Guarantee: ",
        "Every fact is matched to the verified resume — nothing is invented."),
    (10, "Provider-Agnostic Inference: ",
         "Picks the best of 7 LLM backends — free cloud or fully offline."),
    (11, "One-Page ATS-Safe Rendering: ",
         "Single-column PDF, scored by 6 deterministic ATS checks."),
]
for idx, lead, desc in features:
    set_run(paras[idx].runs[0], lead)
    set_run(paras[idx].runs[1], desc, budget=137 - len(lead))

# --- 4. Measurable Metrics bullets ---
set_run(paras[15].runs[0], "Test Coverage: ")
set_run(
    paras[15].runs[1],
    "94 automated tests pass fully offline, enforced on every push via CI.",
    budget=90 - len("Test Coverage: "),
)
set_run(paras[16].runs[0], "Verified Performance: ")
set_run(
    paras[16].runs[1],
    "~70s per run on free inference vs. 1–3 min local; scores of "
    "80–88/100 achieved.",
    budget=97 - len("Verified Performance: "),
)

# --- "Interface Showcase Placeholder" subheading is dropped below (redundant
# with the figure caption, and removing it recovers needed page-1 space) ---

# --- Figure caption ---
set_run(
    paras[18].runs[0],
    "Figure 1: resume-forge web interface — intake form with live "
    "progress tracking.",
    budget=80,
)

# =========================== TABLES ===========================

# --- Table 0: metadata bar ---
meta = d.tables[0]
meta.rows[0].cells[0].paragraphs[0].runs[1].text = " \U0001F7E2 Complete"
meta.rows[0].cells[1].paragraphs[0].runs[1].text = " July 2026"
meta.rows[0].cells[2].paragraphs[0].runs[1].text = " Vishal Poornima"
meta.rows[0].cells[3].paragraphs[0].runs[1].text = " GitHub Repo"

# --- Table 1: Key Impact Metric ---
impact = d.tables[1]
set_run(
    impact.rows[0].cells[0].paragraphs[0].runs[1],
    "One-page, ATS-optimized resume scoring 80+/100 in ~70 seconds.",
    budget=109 - len("\U0001F525 Key Impact Metric: "),
)

# --- Table 2: Technical Implementation (relabeled to the real stack) ---
tech = d.tables[2]
rows_content = [
    ("Frontend", "React, Vite, Tailwind CSS",
     "Animated progress UI, no framework overhead."),
    ("Backend / API", "Python, FastAPI",
     "Background jobs with live progress polling."),
    ("LLM Inference", "Ollama + 6 free/paid cloud presets",
     "Runs unmodified: $0 budget to enterprise Claude."),
    ("Document Rendering", "LaTeX (Tectonic), Jinja2",
     "Deterministic, ATS-safe PDF typesetting."),
]
for i, (layer, techsel, rationale) in enumerate(rows_content, start=1):
    row = tech.rows[i]
    row.cells[0].paragraphs[0].runs[0].text = layer
    row.cells[1].paragraphs[0].runs[0].text = techsel
    row.cells[2].paragraphs[0].runs[0].text = rationale
    budget_report.append((rationale[:40], len(rationale), 55))

# --- Table 3: replace bracket placeholder with the real screenshot ---
shot_cell = d.tables[3].rows[0].cells[0]
shot_para = shot_cell.paragraphs[0]
shot_para.runs[0].text = ""
run = shot_para.add_run()
run.add_picture(SCREENSHOT, width=Inches(1.15))

# Cut the cell's generous top/bottom padding (0.28in each, sized for an empty
# placeholder box) now that it holds a real image -- and drop the "Interface
# Showcase" subheading, which is redundant with the figure caption directly
# below it. Together this recovers close to the exact space needed to keep
# everything on one page.
from docx.oxml.ns import qn as _qn
tcPr = shot_cell._tc.get_or_add_tcPr()
tcMar = tcPr.find(_qn('w:tcMar'))
for side in ("top", "bottom"):
    el = tcMar.find(_qn(f'w:{side}'))
    el.set(_qn('w:w'), "20")

# shave the caption's own "space before" too -- one more small, cheap cut
caption_sp = paras[18]._p.find(_qn('w:pPr')).find(_qn('w:spacing'))
caption_sp.set(_qn('w:before'), "0")

interface_heading = paras[17]._p
interface_heading.getparent().remove(interface_heading)

# Trim the four numbered-section headings' before/after spacing slightly --
# a small, evenly-distributed cut across the whole page rather than one big
# cut in one place.
for idx in (3, 8, 12, 13):
    sp = paras[idx]._p.find(_qn('w:pPr')).find(_qn('w:spacing'))
    sp.set(_qn('w:before'), "200")
    sp.set(_qn('w:after'), "80")

# --- Diagnostic fix: the template glues "3. Technical Implementation" to the
# tech-stack table via keepNext, and every table row is cantSplit + a repeating
# header row. With real (longer-than-placeholder) content earlier on the page,
# some renderers treat heading+table as one oversized atomic block and jump the
# WHOLE thing to page 2 even when ample space remains -- rather than letting the
# heading start where it naturally falls. Removing keepNext here lets the
# heading flow normally; the table's own cantSplit still keeps its rows intact.
from docx.oxml.ns import qn as _qn
heading3_pPr = paras[12]._p.find(_qn('w:pPr'))
kn = heading3_pPr.find(_qn('w:keepNext'))
if kn is not None:
    heading3_pPr.remove(kn)

d.save(OUT)

print("Saved:", OUT)
print("\n=== Length budget check (text, len, budget) ===")
over = 0
for text, length, budget in budget_report:
    flag = "  OVER" if length > budget else "ok"
    if length > budget:
        over += 1
    print(f"[{flag:>6}] {length:>3}/{budget:<3}  {text!r}")
print(f"\n{over} field(s) exceed the template's original one-page budget.")
