"""Step 1: parse the master resume ONCE into a structured profile, with caching."""

from __future__ import annotations

import hashlib
import os
import re
from pathlib import Path

from .exceptions import IngestError
from .llm import LLM, default_llm
from .models import MasterProfile

# Bump when the parsing prompt or schema changes so stale caches are ignored.
PARSE_VERSION = "3"

INGEST_SYSTEM = """You are an expert resume parser. Extract the candidate's resume into the \
provided schema, faithfully and completely.

Rules:
- Copy facts exactly as written: employers, titles, dates, degrees, GPAs, certifications, metrics.
- Do NOT invent, embellish, or normalize away information. If a field is absent, leave it null/empty.
- Keep every bullet point, preserving numbers and technologies mentioned.
- Group skills into sensible categories if the resume does not already group them."""


def default_cache_dir() -> Path:
    env = os.environ.get("RESUME_FORGE_CACHE_DIR")
    if env:
        return Path(env).expanduser()
    return Path.home() / ".cache" / "resume_forge"


def extract_resume_text(path: str | Path) -> str:
    """Extract plain text from a PDF/docx/tex/txt/md resume file."""
    path = Path(path)
    if not path.exists():
        raise IngestError(f"Resume file not found: {path}")
    suffix = path.suffix.lower()

    if suffix == ".pdf":
        import pdfplumber

        with pdfplumber.open(path) as pdf:
            text = "\n".join(page.extract_text() or "" for page in pdf.pages)
    elif suffix == ".docx":
        import docx

        document = docx.Document(str(path))
        parts = [p.text for p in document.paragraphs]
        for table in document.tables:
            for row in table.rows:
                parts.append(" | ".join(cell.text for cell in row.cells))
        text = "\n".join(parts)
    elif suffix in (".tex", ".txt", ".md"):
        text = path.read_text(encoding="utf-8", errors="replace")
    else:
        raise IngestError(
            f"Unsupported resume format '{suffix}'. Use .pdf, .docx, .tex, .txt, or .md."
        )

    if not text.strip():
        raise IngestError(f"No text could be extracted from {path}.")
    return text


_LINKEDIN_RE = re.compile(r"(?:www\.)?linkedin\.com/in/[\w-]+", re.IGNORECASE)
_GITHUB_RE = re.compile(r"(?:www\.)?github\.com/[\w-]+", re.IGNORECASE)


def _looks_gutted(profile: MasterProfile, text: str) -> bool:
    """True when a non-trivial resume parsed into no substance at all."""
    return (
        len(text) > 600
        and not profile.experience
        and not profile.projects
        and not profile.skills
    )


def _backfill_contact(profile: MasterProfile, text: str) -> MasterProfile:
    """Deterministically recover contact fields a (small) model may have dropped.

    Contact info is regex-detectable, so we never rely on the LLM alone for it.
    """
    from .ats import EMAIL_RE, PHONE_RE

    contact = profile.contact
    if not contact.email and (match := EMAIL_RE.search(text)):
        contact.email = match.group(0)
    if not contact.phone and (match := PHONE_RE.search(text)):
        contact.phone = match.group(0).strip()
    if not contact.linkedin and (match := _LINKEDIN_RE.search(text)):
        contact.linkedin = match.group(0)
    if not contact.github and (match := _GITHUB_RE.search(text)):
        contact.github = match.group(0)
    return profile


def _cache_key(path: Path) -> str:
    digest = hashlib.sha256()
    digest.update(PARSE_VERSION.encode())
    digest.update(path.read_bytes())
    return digest.hexdigest()


def ingest_master_resume(
    path: str | Path,
    *,
    llm: LLM | None = None,
    cache_dir: str | Path | None = None,
    use_cache: bool = True,
) -> MasterProfile:
    """Parse the master resume into a :class:`MasterProfile`, cached by file content."""
    path = Path(path)
    cache_dir = Path(cache_dir) if cache_dir else default_cache_dir()
    cache_file = cache_dir / f"profile-{_cache_key(path)}.json"

    if use_cache and cache_file.exists():
        try:
            return MasterProfile.model_validate_json(cache_file.read_text(encoding="utf-8"))
        except Exception:
            pass  # corrupt/stale cache — re-parse

    text = extract_resume_text(path)
    llm = llm or default_llm()
    profile = llm.parse(
        system=INGEST_SYSTEM,
        prompt=f"Parse this resume into the schema:\n\n<resume>\n{text}\n</resume>",
        output_type=MasterProfile,
    )

    # Small models sometimes return schema-valid but EMPTY output (every field
    # has a default, so `{}` validates). Never accept a gutted parse silently:
    # retry, then escalate to the strongest installed model, then error loudly.
    if _looks_gutted(profile, text):
        profile = llm.parse(
            system=INGEST_SYSTEM,
            prompt=(
                "Your previous parse of this resume was empty, which is wrong — the resume "
                "clearly contains skills, experience, and/or projects sections. Parse it "
                "again and extract EVERY entry and EVERY bullet completely.\n\n"
                f"<resume>\n{text}\n</resume>"
            ),
            output_type=MasterProfile,
        )
    if _looks_gutted(profile, text):
        from .llm import stronger_llm_for

        stronger = stronger_llm_for(llm)
        if stronger is not None:
            profile = stronger.parse(
                system=INGEST_SYSTEM,
                prompt=f"Parse this resume into the schema:\n\n<resume>\n{text}\n</resume>",
                output_type=MasterProfile,
            )
    if _looks_gutted(profile, text):
        raise IngestError(
            "The model could not extract skills/experience/projects from this resume "
            "even though it contains substantial text. Try a stronger model "
            "(e.g. `ollama pull qwen2.5:7b` and set RESUME_FORGE_INGEST_MODEL=qwen2.5:7b), "
            "or convert the resume to .txt and retry."
        )

    profile = _backfill_contact(profile, text)

    cache_dir.mkdir(parents=True, exist_ok=True)
    cache_file.write_text(profile.model_dump_json(indent=2), encoding="utf-8")
    return profile
